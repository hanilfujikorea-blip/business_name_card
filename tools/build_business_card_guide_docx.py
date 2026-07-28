"""Build the K Group business-card automation operating guide as a DOCX.

The Markdown source remains the content authority.  This intentionally handles
only the small, documented Markdown subset used by this guide.
"""

from __future__ import annotations

import argparse
import re
import xml.etree.ElementTree as ET
from datetime import datetime, timezone
from zipfile import ZIP_DEFLATED, ZipFile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


FONT = "Malgun Gothic"
INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TABLE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
CONTENT_DXA = 9638
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def xml(tag: str, **attrs: str) -> OxmlElement:
    element = OxmlElement(tag)
    for key, value in attrs.items():
        element.set(qn(key), str(value))
    return element


def set_rfonts(target, name: str = FONT) -> None:
    """Apply Latin and East Asian fonts to a run or style font element."""
    target.font.name = name
    rpr = target._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = xml("w:rFonts")
        rpr.insert(0, fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attr), name)


def shade(cell_or_paragraph, fill: str) -> None:
    props = cell_or_paragraph._tc.get_or_add_tcPr() if hasattr(cell_or_paragraph, "_tc") else cell_or_paragraph._p.get_or_add_pPr()
    props.append(xml("w:shd", **{"w:fill": fill, "w:val": "clear"}))


def cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = xml("w:tcMar")
        tc_pr.append(mar)
    for side, value in CELL_MARGINS.items():
        side_el = mar.find(qn(f"w:{side}"))
        if side_el is None:
            side_el = xml(f"w:{side}")
            mar.append(side_el)
        side_el.set(qn("w:w"), str(value))
        side_el.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = xml("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = xml("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Mm(width / 56.6929)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def mark_header_row(row) -> None:
    row._tr.get_or_add_trPr().append(xml("w:tblHeader", **{"w:val": "true"}))


def add_left_border(paragraph, color: str = BLUE) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = xml("w:pBdr")
        ppr.append(borders)
    borders.append(xml("w:left", **{"w:val": "single", "w:sz": "16", "w:space": "8", "w:color": color}))


def add_bottom_rule(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    borders = xml("w:pBdr")
    borders.append(xml("w:bottom", **{"w:val": "single", "w:sz": "6", "w:space": "6", "w:color": "C9D6E3"}))
    ppr.append(borders)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    run._r.append(xml("w:fldChar", **{"w:fldCharType": "begin"}))
    run._r.append(xml("w:instrText", **{"xml:space": "preserve"}))
    run._r[-1].text = " PAGE "
    run._r.append(xml("w:fldChar", **{"w:fldCharType": "end"}))


def make_numbering(doc, abstract_id: int, num_id: int, fmt: str, marker: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract = xml("w:abstractNum", **{"w:abstractNumId": str(abstract_id)})
    abstract.append(xml("w:multiLevelType", **{"w:val": "singleLevel"}))
    level = xml("w:lvl", **{"w:ilvl": "0"})
    level.append(xml("w:start", **{"w:val": "1"}))
    level.append(xml("w:numFmt", **{"w:val": fmt}))
    level.append(xml("w:lvlText", **{"w:val": marker}))
    level.append(xml("w:lvlJc", **{"w:val": "left"}))
    ppr = xml("w:pPr")
    tabs = xml("w:tabs")
    tabs.append(xml("w:tab", **{"w:val": "num", "w:pos": "540"}))
    ppr.append(tabs)
    ppr.append(xml("w:ind", **{"w:left": "540", "w:hanging": "270"}))
    ppr.append(xml("w:spacing", **{"w:after": "80", "w:line": "300", "w:lineRule": "auto"}))
    level.append(ppr)
    rpr = xml("w:rPr")
    rpr.append(xml("w:rFonts", **{"w:ascii": FONT, "w:hAnsi": FONT, "w:eastAsia": FONT, "w:cs": FONT}))
    level.append(rpr)
    abstract.append(level)
    numbering.append(abstract)
    num = xml("w:num", **{"w:numId": str(num_id)})
    num.append(xml("w:abstractNumId", **{"w:val": str(abstract_id)}))
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = xml("w:numPr")
    num_pr.append(xml("w:ilvl", **{"w:val": "0"}))
    num_pr.append(xml("w:numId", **{"w:val": str(num_id)}))
    ppr.append(num_pr)


def set_style(style, size: float, color: str | None = None, bold: bool | None = None,
              before: float = 0, after: float = 6, line: float = 1.25) -> None:
    set_rfonts(style)
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line


def clear_style_residue(style, *, remove_italic: bool = False) -> None:
    ppr = style._element.pPr
    if ppr is not None:
        for name in ('pBdr', 'numPr'):
            element = ppr.find(qn(f'w:{name}'))
            if element is not None:
                ppr.remove(element)
    if remove_italic and style._element.rPr is not None:
        for name in ('i', 'iCs'):
            element = style._element.rPr.find(qn(f'w:{name}'))
            if element is not None:
                style._element.rPr.remove(element)


def configure_document(doc: Document) -> dict[str, int]:
    section = doc.sections[0]
    section.page_width, section.page_height = Mm(210), Mm(297)
    section.top_margin = section.bottom_margin = Mm(20)
    section.left_margin = section.right_margin = Mm(20)
    section.header_distance = section.footer_distance = Mm(12.5)
    section.different_first_page_header_footer = True

    set_style(doc.styles["Normal"], 10.5, after=6, line=1.25)
    set_style(doc.styles["Title"], 26, INK, True, after=8, line=1.0)
    set_style(doc.styles["Subtitle"], 13, "48647C", False, after=14, line=1.15)
    clear_style_residue(doc.styles['Title'])
    clear_style_residue(doc.styles['Subtitle'], remove_italic=True)
    set_style(doc.styles["Heading 1"], 16, BLUE, True, before=18, after=10, line=1.1)
    set_style(doc.styles["Heading 2"], 13, BLUE, True, before=14, after=7, line=1.1)
    set_style(doc.styles["Heading 3"], 12, DARK_BLUE, True, before=10, after=5, line=1.1)
    set_style(doc.styles["Heading 4"], 11, DARK_BLUE, True, before=8, after=4, line=1.1)
    code = doc.styles.add_style("Guide Code", WD_STYLE_TYPE.PARAGRAPH)
    set_style(code, 9.3, "263745", False, before=4, after=6, line=1.15)
    code.font.name = "Consolas"
    code._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    callout = doc.styles.add_style("Guide Callout", WD_STYLE_TYPE.PARAGRAPH)
    set_style(callout, 10.3, INK, False, before=6, after=8, line=1.25)
    callout.paragraph_format.left_indent = Mm(4)
    callout.paragraph_format.right_indent = Mm(3)
    quote = doc.styles.add_style('Guide Quote', WD_STYLE_TYPE.PARAGRAPH)
    set_style(quote, 10.3, '48647C', False, before=4, after=8, line=1.25)
    quote.paragraph_format.left_indent = Mm(4)
    quote.paragraph_format.right_indent = Mm(3)

    return {
        "bullet": make_numbering(doc, 90, 90, "bullet", "•"),
        "decimal": make_numbering(doc, 91, 91, "decimal", "%1."),
        "check": make_numbering(doc, 92, 92, "bullet", "□"),
    }


def add_runs(paragraph, text: str, code_font: bool = False, bold: bool = False) -> None:
    if code_font:
        run = paragraph.add_run(text)
        run.font.name = 'Consolas'
        run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Consolas')
        run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Consolas')
        return
    parts = re.split(r'(\*\*.+?\*\*|`.+?`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**'):
            add_runs(paragraph, part[2:-2], bold=True)
            continue
        run = paragraph.add_run(part[1:-1] if part.startswith('`') else part)
        run.bold = bold
        if part.startswith('`'):
            run.font.name = 'Consolas'
            run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Consolas')
            run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Consolas')
        else:
            set_rfonts(run)


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(88)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    r = kicker.add_run("K GROUP · INTERNAL REFERENCE GUIDE")
    set_rfonts(r)
    r.font.size, r.font.bold, r.font.color.rgb = Pt(10), True, RGBColor.from_string(BLUE)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(9)
    p.add_run(title)
    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(subtitle)
    doc.add_paragraph().paragraph_format.space_after = Pt(54)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    r = meta.add_run("문서 기준일: 2026-07-27")
    set_rfonts(r)
    r.font.size, r.font.bold, r.font.color.rgb = Pt(10), True, RGBColor.from_string(INK)
    audience = doc.add_paragraph()
    audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = audience.add_run("대상: 운영 담당자 · 관리자 · 인수인계 담당자")
    set_rfonts(r)
    r.font.size, r.font.italic, r.font.color.rgb = Pt(10), True, RGBColor.from_string("5B6B7A")
    doc.add_page_break()


def add_contents(doc: Document, headings: list[str], num_id: int) -> None:
    p = doc.add_paragraph("목차", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    intro = doc.add_paragraph("운영 시 필요한 항목을 빠르게 찾을 수 있도록 구성한 정적 목차입니다.")
    intro.paragraph_format.space_after = Pt(10)
    for heading in headings:
        p = doc.add_paragraph()
        apply_num(p, num_id)
        add_runs(p, re.sub(r"^\d+\.\s*", "", heading))
    doc.add_page_break()


def parse_table(lines: list[str], index: int) -> tuple[list[list[str]], int]:
    rows = []
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [c.strip() for c in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            rows.append(cells)
        index += 1
    return rows, index


def column_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    weights = [max(5, max(len(row[i]) for row in rows)) for i in range(cols)]
    # Limit the visual dominance of a verbose column without making labels tiny.
    weights = [min(weight, 42) for weight in weights]
    raw = [max(950, int(CONTENT_DXA * weight / sum(weights))) for weight in weights]
    raw[-1] += CONTENT_DXA - sum(raw)
    return raw


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    widths = column_widths(rows)
    set_table_geometry(table, widths)
    for row_i, source_row in enumerate(rows):
        for col_i, text in enumerate(source_row):
            cell = table.cell(row_i, col_i)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            add_runs(p, text)
            if row_i == 0:
                shade(cell, TABLE_FILL)
                for run in p.runs:
                    run.bold = True
    mark_header_row(table.rows[0])
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)


def add_callout(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Guide Callout")
    shade(p, CALLOUT_FILL)
    add_left_border(p)
    add_runs(p, text)


def add_quote(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style='Guide Quote')
    add_runs(paragraph, text)


def render_markdown(doc: Document, source: str, numbering: dict[str, int]) -> tuple[int, int, list[str]]:
    lines = source.splitlines()
    headings = []
    table_count = 0
    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph(style="Guide Code")
                shade(p, "EEF2F6")
                add_left_border(p, "9FB5C9")
                add_runs(p, "\n".join(code_lines), code_font=True)
                code_lines = []
            in_code = not in_code
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            table_count += 1
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            level, text = len(heading.group(1)), heading.group(2)
            if level == 1:
                # The sole source H1 is already represented by the cover.
                i += 1
                continue
            style = f"Heading {level - 1}"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.keep_with_next = True
            add_runs(p, text)
            if level == 2:
                headings.append(text)
            i += 1
            continue
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            add_bottom_rule(p)
            i += 1
            continue
        if stripped.startswith(">"):
            quote_text = stripped[1:].strip()
            visible_label = re.sub(r'^[*_`]+', '', quote_text).lstrip()
            if visible_label.startswith(('주의', '중요')):
                add_callout(doc, quote_text)
            else:
                add_quote(doc, quote_text)
            i += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        bullet = re.match(r"^-\s+(.*)$", stripped)
        if numbered:
            p = doc.add_paragraph()
            apply_num(p, numbering["decimal"])
            add_runs(p, numbered.group(1))
            i += 1
            continue
        if bullet:
            text = bullet.group(1)
            is_check = text.startswith("[ ] ")
            p = doc.add_paragraph()
            apply_num(p, numbering["check"] if is_check else numbering["bullet"])
            add_runs(p, text[4:] if is_check else text)
            i += 1
            continue
        if stripped:
            p = doc.add_paragraph()
            add_runs(p, stripped)
        i += 1
    return len(headings), table_count, headings


def configure_furniture(doc: Document) -> None:
    section = doc.sections[0]
    # First page is intentionally clear; subsequent pages carry the quiet label.
    section.first_page_header.paragraphs[0].text = ""
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(3)
    r = header.add_run("K Group 명함 자동발주 · 개발·운영 지침서")
    set_rfonts(r)
    r.font.size, r.font.color.rgb = Pt(8.5), RGBColor.from_string("6B7E90")
    add_bottom_rule(header)
    for footer in (section.footer, section.first_page_footer):
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("K Group  |  ")
        set_rfonts(r)
        r.font.size, r.font.color.rgb = Pt(8.5), RGBColor.from_string("6B7E90")
        add_page_field(p)


def set_extended_company(output_path: Path) -> None:
    """Serialize the Company extended property that python-docx does not expose."""
    temp_path = output_path.with_suffix(".company.tmp.docx")
    company_tag = "{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}Company"
    app_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}'
    word_tag = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
    with ZipFile(output_path, "r") as source, ZipFile(temp_path, "w", ZIP_DEFLATED) as target:
        document_xml = ET.fromstring(source.read('word/document.xml'))
        word_count = sum(len((node.text or '').split()) for node in document_xml.iter(word_tag))
        for info in source.infolist():
            payload = source.read(info.filename)
            if info.filename == "docProps/app.xml":
                app_properties = ET.fromstring(payload)
                company = app_properties.find(company_tag)
                if company is None:
                    company = ET.SubElement(app_properties, company_tag)
                company.text = "K Group"
                pages = app_properties.find(f'{app_ns}Pages')
                if pages is not None:
                    app_properties.remove(pages)
                words = app_properties.find(f'{app_ns}Words')
                if words is None:
                    words = ET.SubElement(app_properties, f'{app_ns}Words')
                words.text = str(word_count)
                payload = ET.tostring(app_properties, encoding="utf-8", xml_declaration=True)
            target.writestr(info, payload)
    temp_path.replace(output_path)

def build(source_path: Path, output_path: Path) -> tuple[int, int]:
    if not source_path.exists():
        raise FileNotFoundError(f"Source Markdown was not found: {source_path}")
    source = source_path.read_text(encoding="utf-8").strip()
    if not source:
        raise ValueError(f"Source Markdown is empty: {source_path}")
    title_match = re.search(r"^#\s+(.+)$", source, re.MULTILINE)
    if not title_match:
        raise ValueError("Source Markdown must begin with a level-one title.")
    title = title_match.group(1)
    section_names = re.findall(r"^##\s+(.+)$", source, re.MULTILINE)
    if len(section_names) != 18:
        raise ValueError(f"Expected 18 top-level sections, found {len(section_names)}.")

    doc = Document()
    properties = doc.core_properties
    basis_date = datetime(2026, 7, 27, tzinfo=timezone.utc)
    properties.author = "K Group"
    properties.last_modified_by = "K Group"
    properties.created = basis_date
    properties.modified = basis_date
    properties.title = title
    properties.subject = "명함 자동발주 시스템 개발·운영 지침서"
    numbering = configure_document(doc)
    configure_furniture(doc)
    add_cover(doc, title, "명함 자동발주 시스템 · 개발 및 운영 기준")
    add_contents(doc, section_names, numbering["decimal"])
    heading_count, table_count, _ = render_markdown(doc, source, numbering)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    set_extended_company(output_path)
    return heading_count, table_count


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    headings, tables = build(args.source, args.output)
    print(f"Output: {args.output.resolve()}")
    print(f"Summary: headings={headings}; tables={tables}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
