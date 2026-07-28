import base64
import re
import tempfile
import unittest
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document

from tools.build_business_card_guide_docx import build


ROOT = Path(__file__).resolve().parents[1]
GUIDE_MD = ROOT / "개발운영지침서_KGroup_명함자동발주.md"
GUIDE_DOCX = ROOT / "개발운영지침서_KGroup_명함자동발주.docx"
LEGACY_MD = ROOT / "개발지침서_명함자동화.md"


EXPECTED_LABEL = '목차'
CANONICAL_SECTIONS = base64.b64decode(
    'MS4g66y47IScIOuqqeyggeqzvCDsi5zsiqTthZwg64+E7J6FIOuwsOqyvQoyLiDsi5zsiqTthZwg7ZWc64iI7JeQIOuztOq4sAozLiDsoITssrQg7JeF66y0IO2dkOumhAo0LiDrjIDsi5zrs7Trk5zsmYAg7Jq07JiBIO2PrO2EuOydmCDsl63tlaAKNS4g7IKs7JqpIOq4sOyIoOqzvCDshKDtg50g7J207JygCjYuIOq1rOyEsSDsmpTshozsmYAg7YyM7J2867OEIOyXre2VoAo3LiDrqZTsnbwg7IiY7KeRLCDslpHsi50g6rKA7KadLCDstIjslYgg7IOd7ISx6rO8IOuwnOyGoSDqs7zsoJUKOC4g642w7J207YSwIOq0gOumrCDrsKnsi53qs7wgREIg7IKs7JqpIOyXrOu2gAo5LiDrqZTsnbztlajqs7wg7JeF7LK0IOyImOyLoMK37LC47KGwIOyjvOyGjOyZgCDrrLjqtawg7ISk7KCVCjEwLiDsp4HsoJEg7Iq57J246rO8IOyekOuPmSDrsJzshqEg66qo65OcCjExLiDsnbzsg4Eg7Jq07JiBIOyInOyEnOyZgCDssrTtgazrpqzsiqTtirgKMTIuIOyYpOulmCDrjIDsnZE6IOycoO2YleqzvCDsobDsuZgg67Cp67KVCjEzLiDrsLHsl4Xqs7wg67O16rWsCjE0LiDqsJzsnbjsoJXrs7TsmYAg67O07JWIIOyjvOydmOyCrO2VrQoxNS4g67OA6rK9IOyLnCDso7zsnZjtlaAg7ISk7KCV6rO8IO2MjOydvAoxNi4g7YWM7Iqk7Yq47JmAIOyatOyYgSDtmZXsnbgg67Cp7IudCjE3LiDsnpDso7wg66y764qUIOyniOusuAoxOC4g7J247IiY7J246rOEIOyytO2BrOumrOyKpO2KuA=='
).decode('utf-8').splitlines()
CANONICAL_SECTIONS[3] = '4. 대시보드와 운영포털의 역할'
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
EP = '{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}'
CP = '{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}'
DCTERMS = '{http://purl.org/dc/terms/}'


class DevelopmentGuideDocumentTests(unittest.TestCase):
    def test_markdown_documents_current_operating_contract(self):
        text = GUIDE_MD.read_text(encoding="utf-8")
        required = [
            "K Group 명함 자동발주 시스템 개발·운영 지침서",
            "별도 데이터베이스를 사용하지 않습니다",
            "processed_state.json",
            "business_card_drafts.json",
            "직접 승인",
            "자동 발송",
            "recover_business_card_state.py",
            "메일 **제목에만** 정확한 부분문자열 `명함`이 들어 있는 경우입니다",
            "첨부 파일명은 대상 선정에 사용하지 않습니다",
        ]
        for phrase in required:
            self.assertIn(phrase, text)
        self.assertRegex(text, r"\.xlsx.*\.xlsm")
        self.assertIn("구형 .xls 형식은 지원하지 않습니다", text)

    def test_markdown_does_not_embed_secret_assignments(self):
        text = GUIDE_MD.read_text(encoding="utf-8")
        self.assertIsNone(re.search(
            r"(?im)^\s*(ARCHIVE_PASSWORD|ARCHIVE_USERNAME|[^\n=]*TOKEN[^\n=]*)\s*=\s*\S+",
            text,
        ))

    def test_legacy_guide_points_to_canonical_guide(self):
        text = LEGACY_MD.read_text(encoding="utf-8")
        self.assertIn(GUIDE_MD.name, text)

    def test_word_artifact_contains_title_and_core_sections(self):
        self.assertTrue(GUIDE_DOCX.exists())
        with zipfile.ZipFile(GUIDE_DOCX) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")
        for phrase in (
            "K Group 명함 자동발주 시스템",
            "데이터 관리",
            "오류 대응",
            "인수인계 체크리스트",
        ):
            self.assertIn(phrase, xml)


    def test_word_heading_ones_match_all_canonical_sections_exactly(self):
        markdown = GUIDE_MD.read_text(encoding='utf-8').splitlines()
        markdown_headings = [line[3:] for line in markdown if line.startswith('## ')]
        self.assertEqual(markdown_headings, CANONICAL_SECTIONS)
        document = Document(GUIDE_DOCX)
        headings = [p.text for p in document.paragraphs if p.style.name == 'Heading 1']
        self.assertEqual(headings[1:], CANONICAL_SECTIONS)

    def test_word_visible_text_strips_nested_inline_delimiters(self):
        text = ' '.join(p.text for p in Document(GUIDE_DOCX).paragraphs)
        self.assertNotIn('Q. `.xls` 파일을 받을 수 있나요?', text)
        self.assertIn('Q. .xls 파일을 받을 수 있나요?', text)

    def test_fenced_code_preserves_intentional_backticks(self):
        separator = chr(10)
        source = '# Fixture' + separator * 2
        source += separator.join('## ' + heading for heading in CANONICAL_SECTIONS)
        source += separator * 2 + '```text' + separator
        source += 'literal `keep` markers' + separator + '```'
        with tempfile.TemporaryDirectory() as temp_dir:
            markdown_path = Path(temp_dir) / 'fixture.md'
            docx_path = Path(temp_dir) / 'fixture.docx'
            markdown_path.write_text(source, encoding='utf-8')
            build(markdown_path, docx_path)
            text = ' '.join(p.text for p in Document(docx_path).paragraphs)
        self.assertIn('literal `keep` markers', text)

    def test_title_and_subtitle_styles_have_no_template_residue(self):
        with zipfile.ZipFile(GUIDE_DOCX) as archive:
            styles = ET.fromstring(archive.read('word/styles.xml'))
        title = next(s for s in styles.findall(f'{W}style') if s.get(f'{W}styleId') == 'Title')
        subtitle = next(s for s in styles.findall(f'{W}style') if s.get(f'{W}styleId') == 'Subtitle')
        self.assertIsNone(title.find(f'{W}pPr/{W}pBdr'))
        self.assertEqual(title.find(f'{W}rPr/{W}sz').get(f'{W}val'), '52')
        self.assertEqual(title.find(f'{W}rPr/{W}color').get(f'{W}val'), '0B2545')
        self.assertIsNone(subtitle.find(f'{W}pPr/{W}pBdr'))
        self.assertIsNone(subtitle.find(f'{W}pPr/{W}numPr'))
        self.assertIsNone(subtitle.find(f'{W}rPr/{W}i'))
        self.assertIsNone(subtitle.find(f'{W}rPr/{W}iCs'))
        self.assertEqual(subtitle.find(f'{W}rPr/{W}sz').get(f'{W}val'), '26')

    def test_document_metadata_uses_basis_date_and_honest_counts(self):
        with zipfile.ZipFile(GUIDE_DOCX) as archive:
            core = ET.fromstring(archive.read('docProps/core.xml'))
            app = ET.fromstring(archive.read('docProps/app.xml'))
        created = core.find(f'{DCTERMS}created').text
        modified = core.find(f'{DCTERMS}modified').text
        self.assertTrue(created.startswith('2026-07-27'))
        self.assertTrue(modified.startswith('2026-07-27'))
        pages = app.find(f'{EP}Pages')
        words = app.find(f'{EP}Words')
        self.assertTrue(pages is None or pages.text != '1')
        self.assertIsNotNone(words)
        self.assertGreater(int(words.text), 0)

    def test_unlabeled_blockquote_is_not_rendered_as_callout(self):
        document = Document(GUIDE_DOCX)
        quote = next(p for p in document.paragraphs if p.text.startswith('이 문서는'))
        self.assertEqual(quote.style.name, 'Guide Quote')
        properties = quote._p.pPr
        self.assertIsNone(properties.find(f'{W}shd'))
        self.assertIsNone(properties.find(f'{W}pBdr'))

    def test_word_extended_properties_identify_k_group(self):
        """The distribution DOCX must serialize K Group as its company."""
        self.assertTrue(GUIDE_DOCX.exists())
        with zipfile.ZipFile(GUIDE_DOCX) as archive:
            app_properties = ET.fromstring(archive.read("docProps/app.xml"))
        company = app_properties.find(
            "{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}Company"
        )
        self.assertIsNotNone(company)
        self.assertEqual(company.text, "K Group")
if __name__ == "__main__":
    unittest.main()
